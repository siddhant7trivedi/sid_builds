import io
import re


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Extract text from a resume file (PDF or DOCX)."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif name.endswith(".docx"):
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Upload a PDF or DOCX.")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF — PyMuPDF first, pdfplumber as fallback."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text() for page in doc]
        text = "\n".join(pages)
        if text.strip():
            return _clean_text(text)
    except Exception:
        pass

    import pdfplumber

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return _clean_text("\n".join(pages))


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return _clean_text("\n".join(parts))


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove junk characters."""
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
